from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler:  
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        doc = Document(self.file_path)
        text = doc.text
        return text


    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        doc = Document(self.file_path)
        doc.add_heading(content, 0)
        doc.styles['Heading 1'].font.size = Pt(font_size)
        doc.styles['Heading 1'].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.styles['Heading 1'].font.bold = True
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline_style = WD_UNDERLINE.SINGLE
        doc.styles['Heading 1'].font.underline_thickness = 1
        doc.styles['Heading 1'].font.name = 'Arial'
        doc.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)
        doc.styles['Heading 1'].font.underline = True
        doc.styles['Heading 1'].font.underline_color.