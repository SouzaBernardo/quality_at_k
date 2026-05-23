from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler: 
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        

        document = Document(self.file_path)
        return document.paragraphs[0].text


    def write_text(self, content, font_size=12, alignment='left'):
        

        document = Document(self.file_path)
        paragraph = document.add_paragraph()
        paragraph.alignment = self._get_alignment_value(alignment)
        paragraph.paragraph_format.line_spacing = Pt(font_size)
        paragraph.add_run(content)
        document.save(self.file_path)
        return True


    def add_heading(self, heading, level=1):
        

        document = Document(self.file_path)
        heading_style = document.styles['Heading 1']
        heading_style.font.size = Pt(12)
        heading_style.font.name = 'Calibri'
        heading_style.font.bold = True
        heading_style.font.italic = False
        heading_style.font.underline = False
        heading_style.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.line_spacing = 1
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.page_break_before = False
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.left_indent = 0
        heading_style.paragraph_format.right_indent = 0
        heading_style.paragraph_format.first_line_indent = 0
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.page_break_before = False
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.left_indent = 0
        heading_style.paragraph_format.right_indent = 0
        heading_style.paragraph_format.first_line_indent = 0
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.page_break_before = False
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.left_indent = 0
        heading_style.paragraph_format.right_indent = 0
        heading_style.paragraph_format.first_line_indent = 0
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.page_break_before = False
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.left_indent = 0
        heading_style.paragraph_format.right_indent = 0
        heading_style.paragraph_format.first_line_indent = 0
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.page_break_before = False
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.left_indent = 0
        heading_style.paragraph_format.right_indent = 0
        heading_style.paragraph_format.first_line_indent = 0
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.page_break_before = False
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.left_indent = 0
        heading_style.paragraph_format.right_indent = 0
        heading_style.paragraph_format.first_line_indent = 0
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format.keep_together = True
        heading_style.paragraph_format.page_break_before = False
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        heading_style.paragraph_format.left_indent = 0
        heading_style.paragraph_format.right_indent = 0
        heading_style.paragraph_format.first_line_indent = 0
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        heading_style.paragraph_format.space_before = 0
        heading_style.paragraph_format.space_after = 0
        heading_style.paragraph_format.widow_control = True
        heading_style.paragraph_format


    def add_table(self, data):
        

        document = Document(self.file_path)
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_paragraph('')
        document.add_


    def _get_alignment_value(self, alignment):