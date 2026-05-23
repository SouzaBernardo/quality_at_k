from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler:  
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        
        document = Document(self.file_path)
        return document.paragraphs[0].text


    def write_text(self, content, font_size=12, alignment='left'):
        
        document = Document(self.file_path)
        document.add_paragraph(content, style='List Bullet')
        document.save(self.file_path)
        return True


    def modify_text(self, content, font_size=12, alignment='left'):


    def add_heading(self, heading, level=1):
        

    def add_table(self, table_data, font_size=12, alignment='left'):


    def add_table(self, data):


    def _get_alignment_value(self, alignment):
