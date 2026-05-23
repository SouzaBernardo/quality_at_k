from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler: 
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        document = Document(self.file_path)
        content = ""
        for para in document.paragraphs:
            content += para.text
        return content

    def write_text(self, content, font_size=12, alignment='left'):
        # code to implement the write_text method
        pass

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
    
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        # Add the heading to the document
        new_paragraph = Document().paragraphs.add()
        new_paragraph.text = heading
        new_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT[level]
    
        # Save the document and close it
        doc = Document(self.file_path)
        doc.add_paragraph(new_paragraph)
        doc.save()
        doc.close()
    
        return True

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        # code to add the table to the Word document
        if not self._add_table(data):
            return False
        return True

    def _get_alignment_value(self, alignment):
        alignment = alignment.upper()
        if alignment == 'LEFT':
            return WD_PARAGRAPH_ALIGNMENT_LEFT
        elif alignment == 'CENTER':
            return WD_PARAGRAPH_ALIGNMENT_CENTER
        elif alignment == 'RIGHT':
            return WD_PARAGRAPH_ALIGNMENT_RIGHT
        else:
            raise ValueError('Invalid alignment value: {}'.format(alignment))