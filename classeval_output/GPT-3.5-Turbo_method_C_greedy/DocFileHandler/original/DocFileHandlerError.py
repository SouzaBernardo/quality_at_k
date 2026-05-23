from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler: 
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        doc = Document(self.file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content
    

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        try:
            document = Document(self.file_path)
            paragraph = document.add_paragraph(content)
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            alignment_value = self._get_alignment_value(alignment)
            paragraph.alignment = alignment_value
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error writing text to document: {e}")
            return False
    

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            document = Document(self.file_path)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(heading)
            font = run.font
            font.size = Pt(14 - level * 2)  # Adjust the font size based on the heading level
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Set the alignment to center
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding heading: {e}")
            return False


    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            document = Document(self.file_path)
            table = document.add_table(rows=len(data), cols=len(data[0]))
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    table.cell(i, j).text = str(cell)
            document.save(self.file_path)
            return True
        except Exception as e:
            print(f"Error adding table: {e}")
            return False
    

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError("Invalid alignment value. Must be 'left', 'center', or 'right'.")
    