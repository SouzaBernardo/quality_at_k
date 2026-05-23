from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        document = Document(self.file_path)
        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
        return text

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.alignment = self._get_alignment_value(alignment)
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)
        try:
            document.save(self.file_path)
            return True
        except:
            return False

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        document = Document(self.file_path)
        try:
            document.add_heading(heading, level)
            document.save(self.file_path)
            return True
        except:
            return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        document = Document(self.file_path)
        table = document.add_table(rows=len(data), cols=len(data[0]))
        for i in range(len(data)):
            for j in range(len(data[i])):
                cell = table.cell(i, j)
                cell.text = str(data[i][j])
        document.save(self.file_path)
        return True

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            return WD_PARAGRAPH_ALIGNMENT.LEFT




# create a DocFileHandler object with the path to the Word document file
doc_file = DocFileHandler("example.docx")

# read the content of the Word document
text = doc_file.read_text()
print(text)

# write some text to the Word document
doc_file.write_text("This is some text.", font_size=24, alignment='center')

# add a heading to the Word document
doc_file.add_heading("This is a heading", level=2)

# add a table to the Word document
data = [["Name", "Age", "Gender"], ["Alice", 25, "Female"], ["Bob", 30, "Male"], ["Charlie", 35, "Male"]]
doc_file.add_table(data)