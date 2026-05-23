from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler: 
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        content = self.file_path.split('\n')[0]
        if len(content) > 1:
            content =''.join(content)
        return content

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        # Add some code here to write the content to the Word document
        #...
        return True

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        if level < 3:
            raise ValueError("Level 1 is not a valid level for headings.")
        if not self.has_heading(heading):
            self.create_heading(heading, level)
            return True
        self._write_paragraph(heading, level, Pt(0, 0))
        return False

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        # Check if the data is a list of lists
        if not isinstance(data, list):
            raise ValueError("data must be a list of lists")

        # Create a new Word document
        doc = Document()

        # Add a table header to the document
        table_header = doc.add_paragraph(
             Pt(0, 0, 100, 100),
             alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
        )

        # Iterate through the data and add rows to the table
        for row in data:
            row_text = []
            for cell in row:
                row_text.append(cell.text)
            table_row = doc.add_paragraph(
                 Pt(0, 0, 100, 100),
                 alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
            )
            table_row.text = row_text

        # Save the Word document
        doc.save(self.file_path)
        return True

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return 1
        elif alignment == 'center':
            return 2
        elif alignment == 'right':
            return 3
        else:
            raise ValueError(f"Invalid alignment value: {alignment}")