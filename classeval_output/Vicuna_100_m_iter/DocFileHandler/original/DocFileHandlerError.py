from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler:  
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        document = Document(self.file_path)
        content = ""
        for para in document.paragraphs:
            content += para.text
        return content

    def write_text(self, content, font_size=12, alignment='left'):
        document = Document(self.file_path)
        content = str(content)
        for para in document.paragraphs:
            para.text = content
        document.save()
        return True

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
    
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        # code to add heading goes here
        return True

    def add_table(self, data):
        # code to add table goes here
        return True

    def _get_alignment_value(self, alignment):
        alignment = alignment.lower()
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT_LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT_CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT_RIGHT
        else:
            return 0