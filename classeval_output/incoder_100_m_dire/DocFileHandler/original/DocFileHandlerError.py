from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler: 
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        document = Document(self.file_path)
        paragraphs = document.paragraphs
        paragraphs_text = []
        for paragraph in paragraphs:
            paragraphs_text.append(paragraph.text)
        return paragraphs_text


    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        document = Document(self.file_path)
        paragraph = document.add_paragraph()
        paragraph.alignment = self._get_alignment_value(alignment)
        paragraph.font.size = Pt(font_size)
        paragraph.text = content
        document.save(self.file_path)
        return True


    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        paragraph = Document().add_paragraph()
        paragraph.add_run().text = heading
        paragraph.add_run().font.size = Pt(font_size)
        paragraph.add_run().font.bold = True
        paragraph.add_run().font.color.rgb = RGBColor(0, 0, 0)
        paragraph.add_run().font.color.theme_color = ThemeColor(rgb=RGBColor(0, 0, 0))
        paragraph.add_run().font.color.transparency = 0.5
        paragraph.add_run().font.color.transparency_method = 'srgb'
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.enum.color.ColorType.PERCENTAGE
        paragraph.add_run().font.color.transparency_percentage = 50
        paragraph.add_run().font.color.transparency_method = \
            docx.


    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        table = Document().add_table(rows=len(data), cols=len(data[0]))
        for row in data:
            table.rows[0].cells[0].text = row[0]
            table.rows[0].cells[1].text = row[1]
            table.rows[1].cells[0].text = row[2]
            table.rows[1].cells[1].text = row[3]
            table.rows[2].cells[0].text = row[4]
            table.rows[2].cells[1].text = row[5]
            table.rows[3].cells[0].text = row[6]
            table.rows[3].cells[1].text = row[7]
            table.rows[4].cells[0].text = row[8]
            table.rows[4].cells[1].text = row[9]
            table.rows[5].cells[0].text = row[10]
            table.rows[5].cells[1].text = row[11]
            table.rows[6].cells[0].text = row[12]
            table.rows[6].cells[1].text = row[13]
            table.rows[7].cells[0].text = row[14]
            table.rows[7].cells[1].text = row[15]
            table.rows[8].cells[0].text = row[16]
            table.rows[8].cells[1].text = row[17]
            table.rows[9].cells[0].text = row[18]
            table.rows[9].cells[1].text = row[19]
            table.rows[10].cells[0].text = row[20]
            table.rows[10].cells[1].text = row[21]
            table.rows[11].cells[0].text = row[22]
            table.rows[11].cells[1].text = row[23]
            table.rows[12].cells[0].text = row[24]
            table.rows[12].cells[1].text = row[25]
            table.rows[13].cells[0].text = row[26]
            table.rows[13].cells[1].text = row[27]
            table.rows[14].cells[0].text = row[28]
            table.rows[14].cells[1].text = row[29]
            table.rows[15].cells[0].text = row[30]
            table.rows[15].cells[1].text = row[31]
            table.rows[16].cells[0].text = row[32]
            table.rows[16].cells[1].text = row[33]
            table.rows[17].cells[0].text = row[34]
            table.rows[17].cells[1].text = row[35]
            table.rows[18].cells[0].text = row[36]
            table.rows[18].cells[1].text = row[37]
            table.rows[19].cells[0].text = row[38]
            table.rows[19].cells[1].text = row[39]
            table.rows[20].cells[0].text = row[40]
            table.rows[20].cells[1].text = row[41]
            table.rows[21].cells[0].text = row[42]
            table.rows[21].cells[1].text = row[43]
            table.rows[22].cells[0].text = row[44]
            table.rows[22].cells[1].text = row[45]
            table.rows[23].cells[0].text = row[46]
            table.rows[23].cells[1].text = row[47]
            table.rows[24].cells[0].text = row[48]
            table.rows[24].cells[1].text = row[49]
            table.rows[25].cells[0].text = row[50]
            table.rows[25].cells[1].text = row[51]
            table.rows[26].cells[0].text = row[52]
            table.rows[26].cells[1].text = row[53]
            table.rows[27].cells[0].text = row[54]
            table.rows[27].cells[1].text = row[55]
            table.rows[28].cells[0].text = row[56]
            table.rows[28].cells[1].text = row[57]
            table.rows[29].cells[0].text = row[58]
            table.rows[29].cells[1].text = row[59]
            table.rows[30].cells[0].text = row[60]
            table.rows[30].cells[1].text = row[61]
            table.rows[31].cells[0].text = row[62]
            table.rows[31].cells[1].text = row[63]
            table.rows[32].cells[0].text = row[64]
            table.rows[32].cells[1].text = row[65]
            table.rows[33].cells[0].text = row[66]
            table.rows[33].cells[1].text = row[67]
            table.rows[34].cells[0].text = row[68]
            table.rows[34].cells[1].text = row[69]
            table.rows[35].cells[0].text = row[70]
            table.rows[35].cells[1].text = row[71]
            table.rows[36].cells[0].text = row[72]
            table.rows[36].cells[1].text = row[73]
            table.rows[37].cells[0].text = row[74]
            table.rows[37].cells[1].text = row[75]
            table.rows[38].cells[0].text = row[76]
            table.rows[38].cells[1].text = row[77]
            table.rows[39].cells[0].text = row[78]
            table.rows[39].cells[1].text = row[79]
            table.rows[40].cells[0].text = row[80]
            table.rows[40].cells[1].text = row[81]
            table.rows[41].cells[0].text = row[82]
            table.rows[41].cells[1].text = row[83]
            table.rows[42].cells[0].text = row[84]
            table.rows[42].cells[1].text = row[85]
            table.rows[43].cells[0].text = row[86]
            table.rows[43].cells[1].text = row[87]
            table.rows[44].cells[0].text = row[88]
            table.rows[44].cells[1].text = row[89]
            table.rows[45].cells[0].text = row[90]
            table.rows[45].cells[1].text = row[91]
            table.rows[46].cells[0].text = row[92]
            table.rows[46].cells[1].text = row[93]
            table.rows[47].cells[0].text = row[94]
            table.rows[47].cells[1].text = row[95]
            table.rows[48].cells[0].text = row[96]
            table.rows[48].cells[1].text = row[97]
            table.rows[49].cells[0].text = row[98]
            table.rows[49].cells[1].text = row[99]
            table.rows[50].cells[0].text = row[100]
            table.rows[50].cells[1].text = row[101]
            table.rows[51].cells[0].text = row[102]
            table.rows[51].cells[1].text = row[103]
            table.rows[52].cells[0].text = row[104]
            table.rows[52].cells[1].text = row[105]
            table.rows[53].cells[0].text = row[106]
            table.rows[53].cells[1].text = row[107]
            table.rows[54].cells[0].text = row[108]
            table.rows[54].cells[1].text = row[109]
            table.rows[55].cells[0].text = row[110]
            table.rows[55].cells[1].text = row[111]
            table.rows[56].cells[0].text = row[112]
            table.rows[56].cells[1].text = row[113]
            table.rows[57].cells[0].text = row[114]
            table.rows[57].cells[1].text = row[115]
            table.rows[58].cells[0].text = row[116]
            table.rows[58].cells[1].text = row[117]
            table.rows[59].cells[0].text = row[118]
            table.rows[59].cells[1].text = row[119]
            table.rows[60].cells[0].text = row[120]
            table.rows[60].cells[1].text = row[121]
            table.rows[61].cells[0].text = row[122]
            table.rows[61].cells[1].text = row[123]
            table.rows[62].cells[0].text = row[124]
            table.rows[62].cells[1].text = row[125]
            table.rows[63].cells[0].text = row[126]
            table.rows[63].cells[1].text = row[127]
            table.rows[64].cells[0].text = row[128]
            table.rows[64].cells[1].text = row[129]
            table.rows[65].cells[0].text = row[130]
            table.rows[65].cells[1].text = row[131]
            table.rows[66].cells[0].text = row[132]
            table.rows[66].cells[1].text = row[133]
            table.rows[67].cells[0].text = row[134]
            table.rows[67].cells[1].text = row[135]
            table.rows[68].cells[0].text = row[136]
            table.rows[68].cells[1].text = row[137]
            table.rows[69].cells[0].text = row[138]
            table.rows[69].cells[1].text = row[139]
            table.rows[70].cells[0].text = row[140]
            table.rows[70].cells[1].text = row[141]
            table.rows[71].cells[0].text = row[142]
            table.rows[71].cells[1].text = row[14


    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        alignment = alignment.lower()
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError('Alignment value not recognized.')
