from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    """
    This is a class that handles Word documents and provides functionalities for reading, writing, and modifying the content of Word documents.
    """

    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path

    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        return self.document.text

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        if not self.document.text:
            return False
        content = self.document.text.replace('\n','')
        content = content.replace('\r','')
        content = self._convert_paragraph_align(content, font_size, alignment)
        self.document.text = content
        return True

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        if not self.document.paragraphs:
            return False
        if level < 1 or level > len(self.document.paragraphs) - 1:
            return False
        paragraph = self.document.paragraphs[level - 1]
        if not paragraph.text:
            return False
        if not paragraph.head:
            return False
        if not paragraph.head.text:
            return False
        if not paragraph.head.text.replace('\n','') or \
            not paragraph.head.text.replace('\r',''):
            return False
        if not self._convert_paragraph_align(paragraph.head.text, font_size, alignment):
            return False
        paragraph.head.text = heading
        return True

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        if not self.document.tables:
            return False
        table = self.document.tables[0]
        if not table.text:
            return False
        if not table.text.replace('\n','') or \
            not table.text.replace('\r',''):
            return False
        if not self._convert_paragraph_align(table.text, font_size, alignment):
            return False
        data = list(map(self._convert_paragraph_align, data))
        table.text = data
        return True

    def _convert_paragraph_align(self, text, font_size, alignment):
        """
        Converts the specified paragraph alignment to a string.
        :param text: str, the paragraph text.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment string ('left', 'center', or 'right'; default is 'left').
        :return: str, the paragraph alignment string.
        """
        if not text:
            return None
        if not text.replace('\n',''):
            return None
        if not text.replace('\r',''):
            return None
        align_value = WD_PARAGRAPH_ALIGNMENT.get(align)
        if align_value is None:
            return None
        return align_value.string