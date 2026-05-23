from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler: 
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path



    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        document = Document(self.file_path)
        content = ""
        for paragraph in document.paragraphs:
            content += paragraph.text + "\n"
        return content

    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        document = Document(self.file_path)
        paragraph = document.add_paragraph()
        paragraph.alignment = self._get_alignment_value(alignment)
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)
        document.save(self.file_path)
        return True

    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        document = Document(self.file_path)
        heading_style = document.styles['Heading %d']
        heading_style.font.size = Pt(level * 16)
        heading_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading_paragraph = document.add_paragraph('')
        heading_paragraph.add_run(heading).bold = True
        heading_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.save(self.file_path)
        return True

    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        document = Document(self.file_path)
        table = document.add_table(rows=1, cols=len(data[0]))
        table.style = 'TableGrid'
        for row in data:
            row_cells = table.add_row().cells
            for cell, value in zip(row_cells, row):
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run(value)
                font = run.font
                font.size = Pt(font_size)
                font.bold = True
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        document.save(self.file_path)
        return True

    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        if alignment == 'left':
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        elif alignment == 'center':
            return WD_PARAGRAPH_ALIGNMENT.CENTER
        elif alignment == 'right':
            return WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            raise ValueError('Invalid alignment value.')