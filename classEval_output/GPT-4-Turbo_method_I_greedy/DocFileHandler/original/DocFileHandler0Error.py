from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
class DocFileHandler: 
    def __init__(self, file_path):
        """
        Initializes the DocFileHandler object with the specified file path.
        :param file_path: str, the path to the Word document file.
        """
        self.file_path = file_path


    def read_text(self):
        """
        Reads the content of a Word document and returns it as a string.
        :return: str, the content of the Word document.
        """
        doc = Document(self.file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return ' '.join(full_text)
    def write_text(self, content, font_size=12, alignment='left'):
        """
        Writes the specified content to a Word document.
        :param content: str, the text content to write.
        :param font_size: int, optional, the font size of the text (default is 12).
        :param alignment: str, optional, the alignment of the text ('left', 'center', or 'right'; default is 'left').
        :return: bool, True if the write operation is successful, False otherwise.
        """
        doc = Document()
                para = doc.add_paragraph()
                run = para.add_run(content)
                run.font.size = Pt(font_size)
        
                if alignment == 'left':
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif alignment == 'center':
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alignment == 'right':
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
                try:
                    doc.save(self.file_path)
                    return True
                except Exception as e:
                    print(f"An error occurred: {e}")
                    return False
    def add_heading(self, heading, level=1):
        """
        Adds a heading to the Word document.
        :param heading: str, the text of the heading.
        :param level: int, optional, the level of the heading (1, 2, 3, etc.; default is 1).
        :return: bool, True if the heading is successfully added, False otherwise.
        """
        try:
            doc = Document(self.file_path)
            doc.add_heading(heading, level)
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"An error occurred: {e}")
            return False
    def add_table(self, data):
        """
        Adds a table to the Word document with the specified data.
        :param data: list of lists, the data to populate the table.
        :return: bool, True if the table is successfully added, False otherwise.
        """
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=1, cols=len(data[0]))
    
            for i, item in enumerate(data[0]):
                table.cell(0, i).text = str(item)
    
            for item in data[1:]:
                cells = table.add_row().cells
                for i, element in enumerate(item):
                    cells[i].text = str(element)
    
            doc.save(self.file_path)
            return True
        except Exception as e:
            print(f"An error occurred: {e}")
            return False
    def _get_alignment_value(self, alignment):
        """
        Returns the alignment value corresponding to the given alignment string.
        :param alignment: str, the alignment string ('left', 'center', or 'right').
        :return: int, the alignment value.
        """
        alignment_dict = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_dict.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
